import streamlit as st
import fitz  # PyMuPDF
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
import google.generativeai as genai
from docx import Document
from fpdf import FPDF
from io import BytesIO

# --- CONVERSION FUNCTIONS ---

def create_word_download(text):
    doc = Document()
    doc.add_heading('NexusPath: Life Action Plan', 0)
    doc.add_paragraph(text)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def create_pdf_download(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    # We use multi_cell to handle line breaks and wrapping
    # Note: This basic PDF version works best for Latin-based languages (Eng, Span, etc.)
    pdf.multi_cell(0, 10, txt=text.encode('latin-1', 'replace').decode('latin-1'))
    return pdf.output()

# --- EXTRACTION FUNCTIONS ---

def extract_text_from_pdf(file):
    text = ""
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_epub(file):
    with open("temp.epub", "wb") as f:
        f.write(file.getbuffer())
    book = epub.read_epub("temp.epub")
    text = ""
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_content(), 'html.parser')
        text += soup.get_text() + "\n"
    return text

# --- GEMINI AI LOGIC ---

def generate_gemini_response(content, language, api_key):
    genai.configure(api_key=api_key)
    available_models = []
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                available_models.append(m.name)
    except Exception as e:
        return f"API Error: {e}"

    selected_model = 'models/gemini-1.5-flash' if 'models/gemini-1.5-flash' in available_models else available_models[0]
    
    try:
        model = genai.GenerativeModel(selected_model)
        prompt = f"""
        You are a Master Life Strategist. Synthesize the teachings from ALL these books into one cohesive system.
        OUTPUT LANGUAGE: {language}
        1. INTEGRATED SUMMARY: Overlapping 'Big Ideas'.
        2. CONFLICT RESOLUTION: Balanced 'middle way' if authors disagree.
        3. MASTER ACTION PLAN: Unified 30-day roadmap.
        4. HABIT STACK: 5 daily habits.
        
        BOOKS CONTENT:
        {content[:800000]}
        """
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Failed: {e}"

# --- UI SETUP ---

st.set_page_config(page_title="NexusPath", layout="wide")
st.title("📚 NexusPath: Multi-Book Action Plan")

with st.sidebar:
    api_key = st.text_input("Gemini API Key", type="password")
    output_lang = st.selectbox("Language", ["English", "Spanish", "French", "German", "Portuguese"])
    st.info("Upload multiple books to create a combined strategy.")

uploaded_files = st.file_uploader("Upload Books (PDF/EPUB)", type=["pdf", "epub"], accept_multiple_files=True)

if uploaded_files:
    combined_text = ""
    for f in uploaded_files:
        combined_text += extract_text_from_pdf(f) if f.type == "application/pdf" else extract_text_from_epub(f)
    
    if st.button("Generate Master Plan"):
        if not api_key:
            st.error("API Key missing.")
        else:
            with st.spinner("Synthesizing multiple books into one plan..."):
                result = generate_gemini_response(combined_text, output_lang, api_key)
                st.session_state['result'] = result # Store in session to keep it visible

if 'result' in st.session_state:
    st.markdown("---")
    st.markdown(st.session_state['result'])
    
    # Download Buttons
    st.subheader("📥 Download Your Plan")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.download_button(
            label="Download as Word (.docx)",
            data=create_word_download(st.session_state['result']),
            file_name="Life_Action_Plan.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    with col2:
        # PDF generation can be intensive, so we wrap it
        try:
            pdf_data = create_pdf_download(st.session_state['result'])
            st.download_button(
                label="Download as PDF (.pdf)",
                data=bytes(pdf_data),
                file_name="Life_Action_Plan.pdf",
                mime="application/pdf"
            )
        except:
            st.warning("PDF conversion failed (contains unsupported characters). Use Word.")
            
    with col3:
        st.download_button(
            label="Download as Text (.txt)",
            data=st.session_state['result'],
            file_name="Life_Action_Plan.txt",
            mime="text/plain"
        )
